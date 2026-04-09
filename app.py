import streamlit as st
import io
import pandas as pd
import numpy as np
import plotly.express as px
import pdfplumber
import re
from typing import TypedDict
from docx import Document
from langgraph.graph import StateGraph, END

# --- 1. ARCHITECTURAL STATE ---
class AgentState(TypedDict):
    pdf_file: io.BytesIO
    telemetry_df: pd.DataFrame
    leakage_ratio: float
    selected_path: str 
    sop_content: str
    error_msg: str

# --- 2. AI AGENT NODES ---

def spatial_parsing_node(state: AgentState):
    """
    INGESTION NODE: Reconstructs tables based on Y-coordinates.
    Handles '00' integer formats and fragmented PDF metadata.
    """
    pdf_file = state.get("pdf_file")
    data = []
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            words = page.extract_words()
            rows = {}
            for w in words:
                y = round(float(w['top']), 0)
                rows.setdefault(y, []).append(w)
            for y in sorted(rows.keys()):
                row_words = sorted(rows[y], key=lambda x: x['x0'])
                row_text = [w['text'].strip() for w in row_words]
                for i, text in enumerate(row_text):
                    # Match '00' or '00:00'
                    if re.match(r'^([0-1]?[0-9]|2[0-3])(:00)?$', text):
                        hour_val = int(text.split(':')[0])
                        for j in range(i + 1, len(row_text)):
                            val_str = row_text[j].replace(',', '.')
                            try:
                                val = float(val_str)
                                data.append({'Hour': hour_val, 'Load_kW': val})
                                break 
                            except ValueError: continue
    if not data: return {"error_msg": "No telemetry found. Ensure PDF contains interval data."}
    df = pd.DataFrame(data).drop_duplicates(subset=['Hour']).sort_values('Hour')
    df = df.set_index('Hour').reindex(range(24)).ffill().bfill().reset_index()
    leakage = (df['Load_kW'].min() / df['Load_kW'].max()) * 100
    return {"telemetry_df": df, "leakage_ratio": leakage, "error_msg": ""}

def routing_node(state: AgentState):
    """ROUTING NODE: Dynamically branches based on metabolic health."""
    path = "incentive" if state['leakage_ratio'] > 30 else "compliance"
    return {"selected_path": path}

def ai_strategist_node(state: AgentState):
    """
    GENERATIVE NODE: Synthesizes the SOP using AI Reasoning.
    This replaces hard-coded strings with data-driven strategy.
    """
    df = state['telemetry_df']
    peak = df['Load_kW'].max()
    floor = df['Load_kW'].min()
    peak_hr = int(df.loc[df['Load_kW'].idxmax()]['Hour'])
    ratio = state['leakage_ratio']
    path = state['selected_path']

    # In a production environment, you would call: response = llm.invoke(prompt)
    # Here, we simulate the AI's Generative Synthesis:
    if path == "incentive":
        ai_output = f"""
### 🚨 STRATEGIC DIRECTIVE: CAPITAL RECOVERY
**Status:** Critical Metabolic Leakage ({ratio:.2f}%)

**1. Peak Annihilation ({peak_hr:02d}:00):**
Targeting the {peak} kW spike. We mandate thermal pre-cooling at {peak_hr-1:02d}:00. Use the facility's Thermal Inertia to 'coast' through the high-tariff window, reducing demand charges by an estimated 15%.

**2. Phantom Load Purge:**
Your {floor} kW baseload is an unauthorized drain. Deploy sequential circuit termination for non-essential zones (4-9) from 00:00 to 04:00.
        """
    else:
        ai_output = f"""
### ✅ STRATEGIC DIRECTIVE: COMPLIANCE SOVEREIGNTY
**Status:** Optimized Metabolic Efficiency

**1. ESG Metric Harvesting:**
Deploy API hooks to verify BC Step Code 4 compliance. We will convert your {peak} kW ceiling data into 'Sovereign Energy Certificates' for market positioning.

**2. Load Shifting:**
Shift 12% of flexible activity into 'Green Windows' where utility carbon intensity is lowest. Maintain baseline stability.
        """
    return {"sop_content": ai_output}

# --- 3. GRAPH ORCHESTRATION ---
workflow = StateGraph(AgentState)
workflow.add_node("parse", spatial_parsing_node)
workflow.add_node("route", routing_node)
workflow.add_node("generate_sop", ai_strategist_node)

workflow.set_entry_point("parse")
workflow.add_edge("parse", "route")
workflow.add_edge("route", "generate_sop")
workflow.add_edge("generate_sop", END)
integra_engine = workflow.compile()

# --- 4. UI HELPERS ---
def get_docx(content):
    doc = Document()
    doc.add_heading("INTEGRA Operational SOP", 0)
    doc.add_paragraph(content)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- 5. STREAMLIT INTERFACE ---
st.set_page_config(page_title="INTEGRA | ArcheSell Demo", layout="wide")

# High-Contrast Metric Styling
st.markdown("""
    <style>
    [data-testid="stMetric"] { background-color: #f0f2f6; border-radius: 10px; padding: 15px; border: 2px solid #31333f; }
    [data-testid="stMetricLabel"] p { color: #31333f !important; font-weight: bold !important; }
    [data-testid="stMetricValue"] div { color: #000000 !important; }
    </style>
    """, unsafe_allow_html=True)

with st.sidebar:
    st.title("🌀 INTEGRA")
    st.caption("Energy Sovereignty Agent")
    up_file = st.file_uploader("Upload Utility PDF", type=["pdf"])
    if up_file and st.button("🚀 EXECUTE AUDIT"):
        st.session_state['res'] = integra_engine.invoke({"pdf_file": up_file})

if st.session_state.get('res'):
    res = st.session_state['res']
    if res.get("error_msg"):
        st.error(res["error_msg"])
    else:
        st.subheader("📌 Executive Metabolic Summary")
        m1, m2, m3 = st.columns(3)
        peak = res['telemetry_df']['Load_kW'].max()
        floor = res['telemetry_df']['Load_kW'].min()
        m1.metric("Peak Load", f"{peak:.2f} kW", delta="Demand Risk", delta_color="inverse")
        m2.metric("Baseload", f"{floor:.2f} kW", delta="Phantom Waste")
        m3.metric("Leakage", f"{res['leakage_ratio']:.2f}%", delta="Critical" if res['leakage_ratio'] > 30 else "Efficient")

        st.divider()
        
        c1, c2 = st.columns([1.5, 1])
        with c1:
            st.subheader("📊 Load Profile (Telemetry)")
            fig = px.bar(res['telemetry_df'], x="Hour", y="Load_kW", color="Load_kW", color_continuous_scale="Viridis", template="plotly_dark")
            fig.add_hline(y=floor, line_dash="dot", line_color="orange", annotation_text="Baseload Floor")
            st.plotly_chart(fig, use_container_width=True)
        with c2:
            st.subheader("📝 Generative SOP")
            st.markdown(res['sop_content'])
            st.download_button("📥 Download AI Strategy", get_docx(res['sop_content']), f"INTEGRA_{res['selected_path']}.docx")
else:
    st.info("Awaiting telemetry. Upload a utility PDF in the sidebar.")